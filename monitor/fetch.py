from dataclasses import dataclass
from io import BytesIO
import json
from pathlib import PurePosixPath
import re
from urllib.parse import urljoin, urlparse

from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

from monitor.fingerprint import normalize_text
from monitor.schema import Source


USER_AGENT = "GeologyJobRadar/1.0 (+https://github.com/healerLSC/geology-job-radar)"
TIMEOUT = (10, 30)


@dataclass(frozen=True)
class FetchResult:
    source_id: str
    status: str
    url: str
    text: str
    links: tuple[str, ...]
    content_type: str
    error: str | None = None
    documents: tuple[tuple[str, str], ...] = ()


def create_session() -> requests.Session:
    session = requests.Session()
    session.headers.update({"User-Agent": USER_AGENT, "Accept-Language": "zh-CN,zh;q=0.9"})
    retry = Retry(
        total=2,
        connect=2,
        read=2,
        backoff_factor=0.8,
        status_forcelist=(429, 500, 502, 503, 504),
        allowed_methods=("GET", "POST"),
    )
    session.mount("https://", HTTPAdapter(max_retries=retry))
    session.mount("http://", HTTPAdapter(max_retries=retry))
    return session


def html_to_text(html: str, parser: str = "html.parser") -> str:
    soup = BeautifulSoup(html, parser)
    for element in soup.select("script, style, noscript, nav, footer, svg"):
        element.decompose()
    return normalize_text(soup.get_text(" ", strip=True))


def html_links(html: str, base_url: str, parser: str = "html.parser") -> tuple[str, ...]:
    soup = BeautifulSoup(html, parser)
    links: list[str] = []
    seen: set[str] = set()
    for anchor in soup.find_all("a", href=True):
        href = str(anchor["href"]).strip()
        if not href or href.startswith(("#", "javascript:", "mailto:", "tel:")):
            continue
        absolute = urljoin(base_url, href)
        parsed = urlparse(absolute)
        if parsed.scheme not in {"http", "https"}:
            continue
        clean = parsed._replace(fragment="").geturl()
        if clean not in seen:
            seen.add(clean)
            links.append(clean)
    for element in soup.find_all("link"):
        href = str(element.get("href") or element.get_text(strip=True)).strip()
        if not href:
            continue
        absolute = urljoin(base_url, href)
        parsed = urlparse(absolute)
        if parsed.scheme not in {"http", "https"}:
            continue
        clean = parsed._replace(fragment="").geturl()
        if clean not in seen:
            seen.add(clean)
            links.append(clean)
    return tuple(links)


def extract_document_text(content: bytes, content_type: str, url: str) -> str:
    suffix = PurePosixPath(urlparse(url).path).suffix.lower()
    lowered_type = content_type.lower()
    if suffix == ".pdf" or "application/pdf" in lowered_type:
        reader = PdfReader(BytesIO(content))
        return normalize_text(" ".join(page.extract_text() or "" for page in reader.pages))
    if suffix == ".docx" or "wordprocessingml.document" in lowered_type:
        document = Document(BytesIO(content))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            chunks.extend(cell.text for row in table.rows for cell in row.cells)
        return normalize_text(" ".join(chunks))
    if suffix == ".xlsx" or "spreadsheetml.sheet" in lowered_type:
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        chunks: list[str] = []
        for worksheet in workbook.worksheets:
            for row in worksheet.iter_rows(values_only=True):
                chunks.append(" ".join(str(value) for value in row if value is not None))
        workbook.close()
        return normalize_text(" ".join(chunks))
    if lowered_type.startswith("text/"):
        return normalize_text(content.decode("utf-8", errors="replace"))
    return ""


def _list_text(value: object, key: str | None = None) -> str:
    if not value:
        return "官方未注明"
    if isinstance(value, str):
        return normalize_text(value)
    if isinstance(value, list):
        chunks: list[str] = []
        for item in value:
            if isinstance(item, dict) and key:
                item = item.get(key, "")
            if item:
                chunks.append(str(item))
        return "、".join(dict.fromkeys(chunks)) or "官方未注明"
    return normalize_text(str(value))


def _date_text(value: object) -> str:
    match = re.search(r"\d{4}-\d{2}-\d{2}", str(value or ""))
    return match.group(0) if match else "官方未注明"


def _iguopin_api_documents(source: Source, session: requests.Session) -> FetchResult:
    endpoint = "https://gp-api.iguopin.com/api/jobs/v1/list"
    headers = {
        "Device": "pc",
        "Version": "5.0.0",
        "Subsite": "cujiuye",
        "User-Agent": USER_AGENT,
        "Content-Type": "application/json",
    }
    documents: list[tuple[str, str]] = []
    seen_job_ids: set[str] = set()

    for keyword in source.query_terms:
        for page in range(1, 4):
            response = session.post(
                endpoint,
                json={"page": page, "page_size": 100, "keyword": keyword},
                headers=headers,
                timeout=TIMEOUT,
            )
            response.raise_for_status()
            payload = response.json()
            if payload.get("code") != 200 or not isinstance(payload.get("data"), dict):
                raise ValueError(f"国聘接口返回异常：{payload.get('code')}")
            data = payload["data"]
            rows = data.get("list") or []
            for row in rows:
                if not isinstance(row, dict):
                    continue
                searchable = json.dumps(row, ensure_ascii=False)
                targets_2027 = re.search(r"2027\s*(?:届|年(?:应届|毕业))", searchable)
                recruitment_context = re.search(r"校园招聘|高校毕业生|应届|暑期实习|实习转正", searchable)
                if not targets_2027 or not recruitment_context:
                    continue
                job_id = str(row.get("job_id") or "").strip()
                if not job_id or job_id in seen_job_ids:
                    continue
                seen_job_ids.add(job_id)
                company = normalize_text(str(row.get("company_name") or "招聘单位未注明"))
                recruitment_type = normalize_text(str(row.get("recruitment_type_cn") or "校园招聘"))
                title = f"{company}2027届{recruitment_type}"
                contents = html_to_text(str(row.get("contents") or ""))
                text = "\n".join(
                    (
                        title,
                        f"发布日期：{_date_text(row.get('start_time'))}",
                        f"报名截止时间：{_date_text(row.get('end_time'))}",
                        f"招聘岗位：{normalize_text(str(row.get('job_name') or '官方未注明'))}",
                        f"学历要求：{_list_text(row.get('education_cn'))}",
                        f"专业要求：{_list_text(row.get('major_cn'))}",
                        f"工作地点：{_list_text(row.get('district_list'), 'area_cn')}",
                        f"岗位详情：{contents}" if contents else "岗位详情：官方未注明",
                    )
                )
                documents.append((f"https://cujiuye.iguopin.com/job/detail?id={job_id}", text))

            total = int(data.get("total") or len(rows))
            if page * 100 >= total or not rows:
                break

    return FetchResult(
        source.source_id,
        "success",
        endpoint,
        "国资央企招聘平台公开岗位接口",
        (),
        "application/json",
        None,
        tuple(documents),
    )


def fetch_source(source: Source, session: requests.Session | None) -> FetchResult:
    if source.mode == "restricted":
        return FetchResult(source.source_id, "restricted", source.url, "", (), "", "公开抓取受限")
    if session is None:
        return FetchResult(source.source_id, "failed", source.url, "", (), "", "缺少网络会话")
    try:
        if source.mode == "iguopin-api":
            return _iguopin_api_documents(source, session)
        response = session.get(source.url, timeout=TIMEOUT, allow_redirects=True)
        response.raise_for_status()
        content_type = response.headers.get("content-type", "").split(";", 1)[0].strip().lower()
        suffix = PurePosixPath(urlparse(response.url).path).suffix.lower()
        is_markup = (
            "html" in content_type
            or "xml" in content_type
            or "rss" in content_type
            or suffix in {".html", ".htm", ".xml", ".rss"}
            or not content_type
        )
        if is_markup:
            html = response.text
            parser = "xml" if "xml" in content_type or "rss" in content_type or suffix in {".xml", ".rss"} else "html.parser"
            text = html_to_text(html, parser)
            links = html_links(html, response.url, parser)
        else:
            text = extract_document_text(response.content, content_type, response.url)
            links = ()
        return FetchResult(
            source.source_id,
            "success" if text else "failed",
            response.url,
            text,
            links,
            content_type,
            None if text else "未提取到可用正文",
        )
    except requests.HTTPError as exc:
        status_code = exc.response.status_code if exc.response is not None else None
        status = "restricted" if status_code in {401, 403, 429} else "failed"
        return FetchResult(source.source_id, status, source.url, "", (), "", str(exc))
    except (requests.RequestException, OSError, ValueError) as exc:
        return FetchResult(source.source_id, "failed", source.url, "", (), "", str(exc))
